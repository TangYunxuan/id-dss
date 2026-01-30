"""
Export API router - export complete session data.
"""
from typing import List, Dict, Any
from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select
from datetime import datetime
from io import BytesIO
import re

from app.db import get_session
from app.models import (
    Session as SessionModel,
    DesignStep,
    AIRecommendation,
    UserAction,
)
import json
from xml.sax.saxutils import escape as _xml_escape

router = APIRouter(prefix="/export", tags=["export"])


@router.get("/{session_id}")
def export_session(
    session_id: int,
    db: Session = Depends(get_session),
) -> Dict[str, Any]:
    """
    Export complete session data including all steps, recommendations, and actions.
    """
    # Get session
    session = db.get(SessionModel, session_id)
    if not session:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Session with id {session_id} not found",
        )
    
    # Get all steps for session
    steps_statement = select(DesignStep).where(DesignStep.session_id == session_id)
    steps = db.exec(steps_statement).all()
    
    # Build export data
    export_data = {
        "session": {
            "id": session.id,
            "course_title": session.course_title,
            "level": session.level,
            "modality": session.modality,
            "constraints": session.constraints,
            "learning_objectives": session.learning_objectives,
            "created_at": session.created_at.isoformat(),
        },
        "design_steps": [],
        "summary": {
            "total_steps": len(steps),
            "total_recommendations": 0,
            "total_actions": 0,
            "actions_by_type": {},
        },
        "exported_at": datetime.utcnow().isoformat(),
    }
    
    # Process each step
    for step in steps:
        step_data = {
            "id": step.id,
            "phase": step.phase,
            "user_input": step.user_input,
            "created_at": step.created_at.isoformat(),
            "recommendations": [],
            "user_actions": [],
        }
        
        # Get recommendations for this step
        rec_statement = select(AIRecommendation).where(AIRecommendation.step_id == step.id)
        recommendations = db.exec(rec_statement).all()
        
        for rec in recommendations:
            # Parse raw_response if it's JSON
            try:
                parsed_response = json.loads(rec.raw_response)
            except (json.JSONDecodeError, TypeError):
                parsed_response = rec.raw_response
            
            rec_data = {
                "id": rec.id,
                "phase": rec.phase,
                "response": parsed_response,
                "created_at": rec.created_at.isoformat(),
            }
            step_data["recommendations"].append(rec_data)
            export_data["summary"]["total_recommendations"] += 1
        
        # Get user actions for this step
        action_statement = select(UserAction).where(UserAction.step_id == step.id)
        actions = db.exec(action_statement).all()
        
        for action in actions:
            action_data = {
                "id": action.id,
                "recommendation_id": action.recommendation_id,
                "action_type": action.action_type,
                "edited_content": action.edited_content,
                "comment": action.comment,
                "created_at": action.created_at.isoformat(),
            }
            step_data["user_actions"].append(action_data)
            export_data["summary"]["total_actions"] += 1
            
            # Count by action type
            action_type = action.action_type
            if action_type not in export_data["summary"]["actions_by_type"]:
                export_data["summary"]["actions_by_type"][action_type] = 0
            export_data["summary"]["actions_by_type"][action_type] += 1
        
        export_data["design_steps"].append(step_data)
    
    return export_data


def _stringify(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (dict, list)):
        return json.dumps(value, ensure_ascii=False, indent=2)
    return str(value)


def _parse_iso(dt: Any) -> datetime:
    if isinstance(dt, datetime):
        return dt
    s = _stringify(dt).strip()
    if not s:
        return datetime.min
    if s.endswith("Z"):
        s = s[:-1] + "+00:00"
    try:
        return datetime.fromisoformat(s)
    except Exception:
        return datetime.min


def _strip_markdown(text: str) -> str:
    """
    Best-effort markdown cleanup for LLM outputs.
    Keeps plain text, removes common markdown artifacts.
    """
    s = text or ""
    # Code fences
    s = re.sub(r"```[\s\S]*?```", lambda m: re.sub(r"^```[a-zA-Z]*\n?", "", m.group(0)).replace("```", ""), s)
    # Inline code
    s = re.sub(r"`([^`]+)`", r"\1", s)
    # Links: [text](url) -> text
    s = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", s)
    # Headings / blockquotes / list markers
    s = re.sub(r"^\s{0,3}#{1,6}\s+", "", s, flags=re.MULTILINE)
    s = re.sub(r"^\s{0,3}>\s?", "", s, flags=re.MULTILINE)
    s = re.sub(r"^\s*(?:[-*•]|\d+\.)\s+", "", s, flags=re.MULTILINE)
    # Emphasis
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"__([^_]+)__", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"_([^_]+)_", r"\1", s)
    # Collapse trailing whitespace per-line
    s = re.sub(r"[ \t]+$", "", s, flags=re.MULTILINE)
    return s.strip()


def _clean_text(value: Any) -> str:
    return _strip_markdown(_stringify(value)).strip()


def _clean_lines(value: Any) -> List[str]:
    raw = _clean_text(value)
    if not raw:
        return []
    return [ln.strip() for ln in raw.splitlines() if ln.strip()]


def _steps_sorted(export_data: Dict[str, Any]) -> List[Dict[str, Any]]:
    steps = export_data.get("design_steps") or []
    return sorted(steps, key=lambda s: _parse_iso(s.get("created_at")))


def _latest_step(export_data: Dict[str, Any], phase: str) -> Dict[str, Any] | None:
    candidates = [s for s in _steps_sorted(export_data) if _stringify(s.get("phase")) == phase]
    return candidates[-1] if candidates else None


def _latest_recommendation_response(export_data: Dict[str, Any], phase: str) -> Any:
    step = _latest_step(export_data, phase)
    if not step:
        return None
    recs = step.get("recommendations") or []
    if not recs:
        return None
    recs_sorted = sorted(recs, key=lambda r: _parse_iso(r.get("created_at")))
    return recs_sorted[-1].get("response")


def _parse_activity_obj(value: Any) -> Dict[str, Any] | None:
    if isinstance(value, dict):
        return value
    if isinstance(value, str) and value.strip():
        try:
            obj = json.loads(value)
            if isinstance(obj, dict):
                return obj
        except Exception:
            return None
    return None


def _activity_id_from_action(action: Dict[str, Any]) -> str | None:
    obj = _parse_activity_obj(action.get("edited_content"))
    if isinstance(obj, dict) and isinstance(obj.get("id"), str) and obj.get("id"):
        return obj.get("id")
    c = _stringify(action.get("comment")).strip()
    if c.startswith("activity-"):
        return c
    return None


def _final_objective_analysis(export_data: Dict[str, Any]) -> Dict[str, Any] | None:
    resp = _latest_recommendation_response(export_data, "objective-analysis")
    if not isinstance(resp, dict):
        return None
    bloom = resp.get("bloom_analysis")
    bloom_clean: List[Dict[str, Any]] = []
    if isinstance(bloom, list):
        for item in bloom:
            if not isinstance(item, dict):
                continue
            bloom_clean.append(
                {
                    "objective": _clean_text(item.get("objective")),
                    "current_level": _clean_text(item.get("current_level")),
                    "domain": _clean_text(item.get("domain")),
                    "is_measurable": bool(item.get("is_measurable")),
                    "suggestion": _clean_text(item.get("suggestion")),
                }
            )
    improved = resp.get("improved_objectives")
    improved_clean = [_clean_text(x) for x in improved] if isinstance(improved, list) else []
    missing = resp.get("missing_coverage")
    missing_clean = [_clean_text(x) for x in missing] if isinstance(missing, list) else []
    return {
        "overall_assessment": _clean_text(resp.get("overall_assessment")),
        "alignment_notes": _clean_text(resp.get("alignment_notes")),
        "bloom_analysis": bloom_clean,
        "improved_objectives": [x for x in improved_clean if x],
        "missing_coverage": [x for x in missing_clean if x],
    }


def _final_activities(export_data: Dict[str, Any]) -> Dict[str, Any] | None:
    step = _latest_step(export_data, "activity-suggestion")
    if not step:
        return None

    resp = _latest_recommendation_response(export_data, "activity-suggestion")
    base_list: List[Dict[str, Any]] = []
    sequence_rationale = ""
    total_estimated_time = ""
    if isinstance(resp, dict):
        base = resp.get("activities") or []
        if isinstance(base, list):
            for i, item in enumerate(base):
                if isinstance(item, dict):
                    base_list.append({"id": f"activity-{i}", **item})
        sequence_rationale = _clean_text(resp.get("sequence_rationale"))
        total_estimated_time = _clean_text(resp.get("total_estimated_time"))

    actions = sorted(step.get("user_actions") or [], key=lambda a: _parse_iso(a.get("created_at")))
    state: Dict[str, Dict[str, Any]] = {}
    for a in actions:
        aid = _activity_id_from_action(a)
        if not aid:
            continue
        action_type = _stringify(a.get("action_type")).strip()
        entry = state.get(aid) or {"status": "pending", "activity": None}
        if action_type == "reject":
            entry["status"] = "rejected"
        elif action_type == "accept":
            entry["status"] = "accepted"
        elif action_type == "edit":
            obj = _parse_activity_obj(a.get("edited_content"))
            if obj:
                entry["activity"] = obj
        state[aid] = entry

    def normalize_activity(a: Dict[str, Any], aid: str, status: str) -> Dict[str, Any]:
        a = dict(a)
        a.pop("draft", None)
        a.pop("status", None)
        a["id"] = aid
        a["status"] = status
        for k in ["title", "type", "description", "duration", "assessment_criteria"]:
            if k in a:
                a[k] = _clean_text(a.get(k))
        for k in ["objective_alignment", "materials_needed", "instructions"]:
            if k in a and isinstance(a.get(k), list):
                a[k] = [_clean_text(x) for x in a.get(k) if _clean_text(x)]
        if isinstance(a.get("adaptations"), dict):
            a["adaptations"] = {kk: _clean_text(vv) for kk, vv in a["adaptations"].items()}
        return a

    merged: List[Dict[str, Any]] = []
    seen: set[str] = set()
    for base in base_list:
        aid = _stringify(base.get("id")).strip()
        if not aid:
            continue
        entry = state.get(aid) or {"status": "pending", "activity": None}
        status = _stringify(entry.get("status") or "pending")
        obj = entry.get("activity") if isinstance(entry.get("activity"), dict) else base
        merged.append(normalize_activity(obj, aid, status))
        seen.add(aid)

    for aid, entry in state.items():
        if aid in seen:
            continue
        obj = entry.get("activity")
        if isinstance(obj, dict):
            merged.append(normalize_activity(obj, aid, _stringify(entry.get("status") or "pending")))

    any_accepted = any(a.get("status") == "accepted" for a in merged)
    if any_accepted:
        final_list = [a for a in merged if a.get("status") == "accepted"]
        selection_note = "Included activities you marked as accepted."
    else:
        final_list = [a for a in merged if a.get("status") != "rejected"]
        selection_note = "No accepted activities were found; included all non-rejected suggestions."

    return {
        "selection_note": selection_note,
        "sequence_rationale": sequence_rationale,
        "total_estimated_time": total_estimated_time,
        "activities": final_list,
    }


def _final_assessments(export_data: Dict[str, Any]) -> Dict[str, Any] | None:
    resp = _latest_recommendation_response(export_data, "assessment-recommendation")
    if not isinstance(resp, dict):
        return None
    assessments = resp.get("assessments")
    if not isinstance(assessments, list):
        return None
    cleaned: List[Dict[str, Any]] = []
    for item in assessments:
        if not isinstance(item, dict):
            continue
        a = dict(item)
        for k in ["title", "type", "method", "description", "timing", "weight", "feedback_strategy"]:
            if k in a:
                a[k] = _clean_text(a.get(k))
        for k in ["objective_alignment", "rubric_criteria"]:
            if k in a and isinstance(a.get(k), list):
                a[k] = [_clean_text(x) for x in a.get(k) if _clean_text(x)]
        cleaned.append(a)
    return {
        "assessment_strategy_rationale": _clean_text(resp.get("assessment_strategy_rationale")),
        "formative_summative_balance": _clean_text(resp.get("formative_summative_balance")),
        "assessments": cleaned,
    }


def _build_final_design(export_data: Dict[str, Any]) -> Dict[str, Any]:
    session = export_data.get("session", {}) or {}
    objectives = _clean_text(session.get("learning_objectives") or "")
    return {
        "title": "Instructional Design Plan",
        "exported_at": _clean_text(export_data.get("exported_at")),
        "session": {
            "id": session.get("id"),
            "course_title": _clean_text(session.get("course_title")),
            "level": _clean_text(session.get("level")),
            "modality": _clean_text(session.get("modality")),
            "constraints": _clean_text(session.get("constraints") or ""),
            "learning_objectives": objectives,
            "learning_objectives_lines": _clean_lines(objectives),
        },
        "objective_analysis": _final_objective_analysis(export_data),
        "activity_plan": _final_activities(export_data),
        "assessment_plan": _final_assessments(export_data),
    }


def _export_to_docx_bytes(export_data: Dict[str, Any]) -> bytes:
    """
    Render a user-friendly final instructional design as a Word document (DOCX).
    Dependency: python-docx
    """
    try:
        from docx import Document  # type: ignore
        from docx.shared import Inches  # type: ignore
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX export requires python-docx. Install backend dependencies and try again.",
        ) from e

    final = _build_final_design(export_data)
    session = final.get("session", {}) or {}
    doc = Document()

    DEFAULT_BULLET_TEXT_INDENT = Inches(0.5)

    def add_bullet(text: str):
        try:
            return doc.add_paragraph(text, style="List Bullet")
        except Exception:
            return doc.add_paragraph(f"- {text}")

    def add_bullet_detail(text: str, anchor_bullet_para=None) -> None:
        """
        Add a paragraph that visually belongs to the bullet above it.
        Indent it to align with the bullet text (not the bullet glyph).
        """
        cleaned = _clean_text(text)
        if not cleaned:
            return
        p = doc.add_paragraph(cleaned)
        indent = None
        if anchor_bullet_para is not None:
            try:
                indent = anchor_bullet_para.paragraph_format.left_indent
            except Exception:
                indent = None
        p.paragraph_format.left_indent = indent or DEFAULT_BULLET_TEXT_INDENT
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0

    doc.add_heading(_clean_text(final.get("title")), level=0)
    doc.add_paragraph(f"Exported at: {_clean_text(final.get('exported_at'))}")
    doc.add_paragraph(f"Session ID: {_stringify(session.get('id'))}")

    doc.add_heading("Course Context", level=1)
    doc.add_paragraph(f"Course Title: {_clean_text(session.get('course_title'))}")
    doc.add_paragraph(f"Level: {_clean_text(session.get('level'))}")
    doc.add_paragraph(f"Modality: {_clean_text(session.get('modality'))}")
    if _clean_text(session.get("constraints")):
        doc.add_heading("Constraints", level=2)
        doc.add_paragraph(_clean_text(session.get("constraints")))

    if session.get("learning_objectives_lines"):
        doc.add_heading("Learning Objectives", level=1)
        for ln in session.get("learning_objectives_lines") or []:
            add_bullet(_clean_text(ln))

    oa = final.get("objective_analysis")
    if isinstance(oa, dict):
        doc.add_heading("Objective Analysis (AI)", level=1)
        if oa.get("overall_assessment"):
            doc.add_paragraph(_clean_text(oa.get("overall_assessment")))
        if oa.get("alignment_notes"):
            doc.add_heading("Alignment Notes", level=2)
            doc.add_paragraph(_clean_text(oa.get("alignment_notes")))
        bloom = oa.get("bloom_analysis") or []
        if bloom:
            doc.add_heading("Bloom's Taxonomy Notes", level=2)
            for item in bloom:
                if not isinstance(item, dict):
                    continue
                line = f"{_clean_text(item.get('objective'))} — level: {_clean_text(item.get('current_level'))}".strip(" —")
                if line:
                    bullet_p = add_bullet(line)
                if item.get("suggestion"):
                    add_bullet_detail(_clean_text(item.get("suggestion")), anchor_bullet_para=bullet_p if line else None)

    ap = final.get("activity_plan")
    if isinstance(ap, dict) and ap.get("activities"):
        doc.add_heading("Learning Activities", level=1)
        if ap.get("selection_note"):
            doc.add_paragraph(_clean_text(ap.get("selection_note")))
        if ap.get("total_estimated_time"):
            doc.add_paragraph(f"Total estimated time: {_clean_text(ap.get('total_estimated_time'))}")
        if ap.get("sequence_rationale"):
            doc.add_heading("Sequence Rationale", level=2)
            doc.add_paragraph(_clean_text(ap.get("sequence_rationale")))

        for idx, act in enumerate(ap.get("activities") or [], start=1):
            if not isinstance(act, dict):
                continue
            title = _clean_text(act.get("title")) or f"Activity {idx}"
            meta = " • ".join([p for p in [_clean_text(act.get("type")), _clean_text(act.get("duration"))] if p])
            doc.add_heading(f"{title}{f' ({meta})' if meta else ''}", level=2)

            if _clean_text(act.get("description")):
                doc.add_paragraph(_clean_text(act.get("description")))
            if isinstance(act.get("objective_alignment"), list) and act.get("objective_alignment"):
                doc.add_paragraph("Aligns with objectives:")
                for o in act.get("objective_alignment"):
                    add_bullet(_clean_text(o))
            if isinstance(act.get("materials_needed"), list) and act.get("materials_needed"):
                doc.add_paragraph("Materials:")
                for m in act.get("materials_needed"):
                    add_bullet(_clean_text(m))
            if isinstance(act.get("instructions"), list) and act.get("instructions"):
                doc.add_paragraph("Instructions:")
                for ins in act.get("instructions"):
                    add_bullet(_clean_text(ins))
            if _clean_text(act.get("assessment_criteria")):
                doc.add_paragraph("Assessment criteria:")
                doc.add_paragraph(_clean_text(act.get("assessment_criteria")))
            if isinstance(act.get("adaptations"), dict) and act.get("adaptations"):
                doc.add_paragraph("Adaptations:")
                for k, v in act.get("adaptations").items():
                    if _clean_text(v):
                        add_bullet(f"{_clean_text(k)}: {_clean_text(v)}")

    asmt = final.get("assessment_plan")
    if isinstance(asmt, dict) and asmt.get("assessments"):
        doc.add_heading("Assessments (AI)", level=1)
        if asmt.get("assessment_strategy_rationale"):
            doc.add_paragraph(_clean_text(asmt.get("assessment_strategy_rationale")))
        if asmt.get("formative_summative_balance"):
            doc.add_paragraph(_clean_text(asmt.get("formative_summative_balance")))

        for idx, a in enumerate(asmt.get("assessments") or [], start=1):
            if not isinstance(a, dict):
                continue
            title = _clean_text(a.get("title")) or f"Assessment {idx}"
            meta = " • ".join([p for p in [_clean_text(a.get("type")), _clean_text(a.get("timing")), _clean_text(a.get("weight"))] if p])
            doc.add_heading(f"{title}{f' ({meta})' if meta else ''}", level=2)
            if _clean_text(a.get("description")):
                doc.add_paragraph(_clean_text(a.get("description")))
            if isinstance(a.get("objective_alignment"), list) and a.get("objective_alignment"):
                doc.add_paragraph("Aligns with objectives:")
                for o in a.get("objective_alignment"):
                    add_bullet(_clean_text(o))
            if isinstance(a.get("rubric_criteria"), list) and a.get("rubric_criteria"):
                doc.add_paragraph("Rubric criteria:")
                for r in a.get("rubric_criteria"):
                    add_bullet(_clean_text(r))
            if _clean_text(a.get("feedback_strategy")):
                doc.add_paragraph("Feedback strategy:")
                doc.add_paragraph(_clean_text(a.get("feedback_strategy")))

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_to_pdf_bytes(export_data: Dict[str, Any]) -> bytes:
    """
    Render a user-friendly final instructional design as a PDF.
    Dependency: reportlab
    """
    try:
        from reportlab.lib.pagesizes import letter  # type: ignore
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak  # type: ignore
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle  # type: ignore
        from reportlab.lib.units import inch  # type: ignore
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF export requires reportlab. Install backend dependencies and try again.",
        ) from e

    def p(text: str) -> str:
        # ReportLab Paragraph supports a small HTML-like subset; escape user content.
        return _xml_escape(text).replace("\n", "<br/>")

    styles = getSampleStyleSheet()
    bullet_style = ParagraphStyle(
        "Bullet",
        parent=styles["Normal"],
        leftIndent=0.5 * inch,        # where bullet text starts
        firstLineIndent=-0.25 * inch, # hang the bullet into the margin
        spaceBefore=0,
        spaceAfter=0,
    )
    bullet_detail_style = ParagraphStyle(
        "BulletDetail",
        parent=styles["Normal"],
        leftIndent=0.5 * inch,  # align with bullet text
        spaceBefore=0,
        spaceAfter=0,
    )

    def add_bullet_pdf(text: str) -> None:
        cleaned = _clean_text(text)
        if not cleaned:
            return
        story.append(Paragraph(p(f"• {cleaned}"), bullet_style))

    def add_bullet_detail_pdf(text: str) -> None:
        cleaned = _clean_text(text)
        if not cleaned:
            return
        story.append(Paragraph(p(cleaned), bullet_detail_style))

    story: List[Any] = []

    final = _build_final_design(export_data)
    session = final.get("session", {}) or {}

    story.append(Paragraph(p(_clean_text(final.get("title"))), styles["Title"]))
    story.append(Paragraph(p(f"Exported at: {_clean_text(final.get('exported_at'))}"), styles["Normal"]))
    story.append(Paragraph(p(f"Session ID: {_stringify(session.get('id'))}"), styles["Normal"]))
    story.append(Spacer(1, 12))

    story.append(Paragraph(p("Course Context"), styles["Heading1"]))
    story.append(Paragraph(p(f"Course Title: {_clean_text(session.get('course_title'))}"), styles["Normal"]))
    story.append(Paragraph(p(f"Level: {_clean_text(session.get('level'))}"), styles["Normal"]))
    story.append(Paragraph(p(f"Modality: {_clean_text(session.get('modality'))}"), styles["Normal"]))
    if _clean_text(session.get("constraints")):
        story.append(Paragraph(p("Constraints"), styles["Heading2"]))
        story.append(Paragraph(p(_clean_text(session.get("constraints"))), styles["Normal"]))
    story.append(Spacer(1, 12))

    if session.get("learning_objectives_lines"):
        story.append(Paragraph(p("Learning Objectives"), styles["Heading1"]))
        for ln in session.get("learning_objectives_lines") or []:
            add_bullet_pdf(_clean_text(ln))
        story.append(Spacer(1, 12))

    oa = final.get("objective_analysis")
    if isinstance(oa, dict):
        story.append(Paragraph(p("Objective Analysis (AI)"), styles["Heading1"]))
        if oa.get("overall_assessment"):
            story.append(Paragraph(p(_clean_text(oa.get("overall_assessment"))), styles["Normal"]))
        if oa.get("alignment_notes"):
            story.append(Paragraph(p("Alignment Notes"), styles["Heading2"]))
            story.append(Paragraph(p(_clean_text(oa.get("alignment_notes"))), styles["Normal"]))
        bloom = oa.get("bloom_analysis") or []
        if bloom:
            story.append(Paragraph(p("Bloom's Taxonomy Notes"), styles["Heading2"]))
            for item in bloom:
                if not isinstance(item, dict):
                    continue
                line = f"{_clean_text(item.get('objective'))} — level: {_clean_text(item.get('current_level'))}".strip(" —")
                if line:
                    add_bullet_pdf(line)
                if item.get("suggestion"):
                    add_bullet_detail_pdf(_clean_text(item.get("suggestion")))
        story.append(Spacer(1, 12))

    ap = final.get("activity_plan")
    if isinstance(ap, dict) and ap.get("activities"):
        story.append(Paragraph(p("Learning Activities"), styles["Heading1"]))
        if ap.get("selection_note"):
            story.append(Paragraph(p(_clean_text(ap.get("selection_note"))), styles["Normal"]))
        if ap.get("total_estimated_time"):
            story.append(Paragraph(p(f"Total estimated time: {_clean_text(ap.get('total_estimated_time'))}"), styles["Normal"]))
        if ap.get("sequence_rationale"):
            story.append(Paragraph(p("Sequence Rationale"), styles["Heading2"]))
            story.append(Paragraph(p(_clean_text(ap.get("sequence_rationale"))), styles["Normal"]))
        story.append(Spacer(1, 6))

        for idx, act in enumerate(ap.get("activities") or [], start=1):
            if not isinstance(act, dict):
                continue
            title = _clean_text(act.get("title")) or f"Activity {idx}"
            meta = " • ".join([p for p in [_clean_text(act.get("type")), _clean_text(act.get("duration"))] if p])
            story.append(Paragraph(p(f"{title}{f' ({meta})' if meta else ''}"), styles["Heading2"]))
            if _clean_text(act.get("description")):
                story.append(Paragraph(p(_clean_text(act.get("description"))), styles["Normal"]))
            if isinstance(act.get("objective_alignment"), list) and act.get("objective_alignment"):
                story.append(Paragraph(p("Aligns with objectives:"), styles["Heading3"]))
                for o in act.get("objective_alignment"):
                    add_bullet_pdf(_clean_text(o))
            if isinstance(act.get("materials_needed"), list) and act.get("materials_needed"):
                story.append(Paragraph(p("Materials:"), styles["Heading3"]))
                for m in act.get("materials_needed"):
                    add_bullet_pdf(_clean_text(m))
            if isinstance(act.get("instructions"), list) and act.get("instructions"):
                story.append(Paragraph(p("Instructions:"), styles["Heading3"]))
                for ins in act.get("instructions"):
                    add_bullet_pdf(_clean_text(ins))
            if _clean_text(act.get("assessment_criteria")):
                story.append(Paragraph(p("Assessment criteria:"), styles["Heading3"]))
                story.append(Paragraph(p(_clean_text(act.get("assessment_criteria"))), styles["Normal"]))
            if isinstance(act.get("adaptations"), dict) and act.get("adaptations"):
                story.append(Paragraph(p("Adaptations:"), styles["Heading3"]))
                for k, v in act.get("adaptations").items():
                    if _clean_text(v):
                        add_bullet_pdf(f"{_clean_text(k)}: {_clean_text(v)}")
            story.append(Spacer(1, 8))

    asmt = final.get("assessment_plan")
    if isinstance(asmt, dict) and asmt.get("assessments"):
        story.append(PageBreak())
        story.append(Paragraph(p("Assessments (AI)"), styles["Heading1"]))
        if asmt.get("assessment_strategy_rationale"):
            story.append(Paragraph(p(_clean_text(asmt.get("assessment_strategy_rationale"))), styles["Normal"]))
        if asmt.get("formative_summative_balance"):
            story.append(Paragraph(p(_clean_text(asmt.get("formative_summative_balance"))), styles["Normal"]))
        for idx, a in enumerate(asmt.get("assessments") or [], start=1):
            if not isinstance(a, dict):
                continue
            title = _clean_text(a.get("title")) or f"Assessment {idx}"
            meta = " • ".join([p for p in [_clean_text(a.get("type")), _clean_text(a.get("timing")), _clean_text(a.get("weight"))] if p])
            story.append(Paragraph(p(f"{title}{f' ({meta})' if meta else ''}"), styles["Heading2"]))
            if _clean_text(a.get("description")):
                story.append(Paragraph(p(_clean_text(a.get("description"))), styles["Normal"]))
            if isinstance(a.get("objective_alignment"), list) and a.get("objective_alignment"):
                story.append(Paragraph(p("Aligns with objectives:"), styles["Heading3"]))
                for o in a.get("objective_alignment"):
                    add_bullet_pdf(_clean_text(o))
            if isinstance(a.get("rubric_criteria"), list) and a.get("rubric_criteria"):
                story.append(Paragraph(p("Rubric criteria:"), styles["Heading3"]))
                for r in a.get("rubric_criteria"):
                    add_bullet_pdf(_clean_text(r))
            if _clean_text(a.get("feedback_strategy")):
                story.append(Paragraph(p("Feedback strategy:"), styles["Heading3"]))
                story.append(Paragraph(p(_clean_text(a.get("feedback_strategy"))), styles["Normal"]))
            story.append(Spacer(1, 8))

    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter, title=_clean_text(final.get("title")))
    doc.build(story)
    return buf.getvalue()


@router.get("/{session_id}/docx")
def export_session_docx(
    session_id: int,
    db: Session = Depends(get_session),
):
    """Export complete session as a Word (DOCX) file."""
    export_data = export_session(session_id=session_id, db=db)
    content = _export_to_docx_bytes(export_data)
    filename = f"id-dss-session-{session_id}.docx"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@router.get("/{session_id}/pdf")
def export_session_pdf(
    session_id: int,
    db: Session = Depends(get_session),
):
    """Export complete session as a PDF file."""
    export_data = export_session(session_id=session_id, db=db)
    content = _export_to_pdf_bytes(export_data)
    filename = f"id-dss-session-{session_id}.pdf"
    headers = {"Content-Disposition": f'attachment; filename="{filename}"'}
    return StreamingResponse(
        BytesIO(content),
        media_type="application/pdf",
        headers=headers,
    )


